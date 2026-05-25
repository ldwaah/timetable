#!/usr/bin/env python3
"""Export ROH role-analysis HTML to professional Word and PDF for HR."""

from __future__ import annotations

import html
import re
import subprocess
import sys
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIRS = [
    SCRIPT_DIR,
    Path("/Users/lloyddwaah/Downloads/Job Description Comparison"),
]

FONT_NAME = "Arial"
FONT_SIZE = Pt(11)
HEADING_COLOR = RGBColor(0x00, 0x00, 0x00)
HEADING_SIZES = {1: Pt(16), 2: Pt(13), 3: Pt(12), 4: Pt(11)}
HIGHLIGHT_FILL = "FFF9CC"  # light yellow
REMOVED_FILL = "F2F2F2"  # subtle grey for removed duties
NOTE_INDENT = Inches(0.35)
SOFFICE = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")

INTRO_REPLACEMENTS = {
    "facilities": (
        "The FC recognises that some elements of the current role overlap with "
        "the proposed role, particularly administrative and coordination tasks. "
        "The concern is not that there is no overlap, but that the proposed "
        "structure changes reporting lines, sign-off, and practical expectations. "
        "Passages highlighted below show where the proposed Estates Support Officer "
        "role differs (ROH describe this as a direct match). For consultation "
        "timing and conclusions reached before practical details were clear, see "
        "Meaningful consultation and ability to assess suitability in Consultation "
        "Response Notes."
    ),
    "consultation": (
        "ROH consultation follow-up letter of 20 May 2026 (meeting of 18 May). "
        "Highlighted passages are disputed or lacked practical clarity at consultation; "
        "notes set out the FC's position on timing, what remained open, and "
        "suitability assessment."
    ),
}


def set_cell_shading(run, fill_hex: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    r_pr.append(shd)


def set_document_defaults(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    for level in range(1, 5):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style.font.size = HEADING_SIZES[level]
        style.font.color.rgb = HEADING_COLOR
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_title_page(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = FONT_NAME
    run.font.color.rgb = HEADING_COLOR
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = FONT_SIZE
        r2.font.name = FONT_NAME
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Prepared for: Facilities Coordinator (FC)")
    r3.font.size = FONT_SIZE
    r3.font.name = FONT_NAME
    doc.add_page_break()


def add_footer(doc: Document, text: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = text
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _normalize_text(line: str) -> str:
    line = html.unescape(line).strip()
    line = line.replace("\u2019", "'").replace("\u2018", "'")
    return re.sub(r"\s+", " ", line)


def _clean_prefix(line: str) -> str:
    line = _normalize_text(line)
    m = re.match(r"^Proposed role:\s*(.+)$", line, re.I)
    if m:
        role = m.group(1).strip().rstrip(".")
        low = role.lower()
        if low in ("estates support officer", "e&f support officer"):
            return "The proposed role is an Estates Support Officer"
        if low.startswith("reports to "):
            return f"The proposed role reports to {role[11:].strip()}"
        if low.startswith("support "):
            return f"In the proposed role, the FC would {role[0].lower()}{role[1:]}"
        return f"In the proposed role: {role}"
    m = re.match(r"^Removed in proposed role\.?\s*(.*)$", line, re.I)
    if m:
        tail = m.group(1).strip()
        base = "This duty is not in the proposed Estates Support Officer role"
        return f"{base}. {tail}" if tail else base
    line = re.sub(r"^FC['']s position \([^)]+\):\s*", "", line, flags=re.I)
    line = re.sub(r"^FC['']s position:\s*", "", line, flags=re.I)
    line = re.sub(r"^Christa['']s position \([^)]+\):\s*", "", line, flags=re.I)
    line = re.sub(r"^Christa['']s position:\s*", "", line, flags=re.I)
    return line.strip()


def format_note(tip_tag: Tag | None) -> str:
    if not tip_tag:
        return ""
    raw = tip_tag.decode_contents()
    raw = re.sub(r"<br\s*/?>", "\n\n", raw, flags=re.I)
    soup = BeautifulSoup(raw, "html.parser")
    text = soup.get_text()
    parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not parts:
        return ""
    cleaned = [_clean_prefix(p) for p in parts]
    if len(cleaned) == 1:
        body = cleaned[0]
    else:
        first, second = cleaned[0], cleaned[1]
        if first.lower().startswith("the proposed estates"):
            body = f"{first.rstrip('.')}. {second}"
        elif first.lower().startswith("this duty is not"):
            body = f"{first.rstrip('.')}. {second}"
        elif first and second:
            if not first.endswith((".", ":", ";")):
                first += "."
            body = f"{first} {second}"
        else:
            body = first or second
    body = re.sub(r"\s+", " ", body).strip()
    if not body.lower().startswith("note"):
        body = f"Note: {body}"
    elif not body.lower().startswith("note:"):
        body = f"Note: {body[4:].lstrip()}"
    if body and body[-1] not in ".!?":
        body += "."
    return body


def add_run(paragraph, text: str, *, bold: bool = False, highlight: str | None = None) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    if highlight:
        set_cell_shading(run, highlight)


def add_note_paragraph(doc: Document, note: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = NOTE_INDENT
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(note)
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def process_inline(
    node: Tag | NavigableString,
    paragraph,
    notes: list[str],
    *,
    removed: bool = False,
) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            add_run(paragraph, text, highlight=REMOVED_FILL if removed else None)
        return

    if not isinstance(node, Tag):
        return

    name = node.name
    classes = node.get("class", [])

    if name == "span" and ("changed" in classes or "removed" in classes):
        is_removed = "removed" in classes
        tip = node.find("span", class_="tip")
        visible = node.get_text()
        if tip:
            visible = visible.replace(tip.get_text(), "")
        add_run(
            paragraph,
            visible,
            highlight=REMOVED_FILL if is_removed else HIGHLIGHT_FILL,
        )
        note = format_note(tip)
        if note:
            notes.append(note)
        return

    if name in ("strong", "b"):
        text = node.get_text()
        add_run(paragraph, text, bold=True, highlight=REMOVED_FILL if removed else None)
        return

    if name == "br":
        add_run(paragraph, "\n")
        return

    for child in node.children:
        process_inline(child, paragraph, notes, removed=removed)


def is_hint(tag: Tag) -> bool:
    return "hint" in tag.get("class", [])


def iter_br_separated_lines(tag: Tag) -> list[str]:
    """Split paragraph content on <br> without leaking '/>' from <br/> markup."""
    lines: list[str] = []
    buf: list[str] = []
    for child in tag.children:
        if isinstance(child, NavigableString):
            buf.append(str(child))
        elif isinstance(child, Tag):
            if child.name == "br":
                line = _normalize_text("".join(buf))
                if line:
                    lines.append(line)
                buf = []
            else:
                buf.append(child.get_text())
    if buf:
        line = _normalize_text("".join(buf))
        if line:
            lines.append(line)
    return lines


def add_block_heading(doc: Document, tag: Tag) -> None:
    text = tag.get_text(strip=True)
    if not text:
        return
    level = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}.get(tag.name, 2)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_NAME
        run.font.size = HEADING_SIZES[level]
        run.font.color.rgb = HEADING_COLOR
        run.bold = True


def add_paragraph_block(doc: Document, tag: Tag, intro_kind: str | None) -> None:
    if is_hint(tag):
        if intro_kind and intro_kind in INTRO_REPLACEMENTS:
            para = doc.add_paragraph(INTRO_REPLACEMENTS[intro_kind])
            para.paragraph_format.space_after = Pt(12)
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                run.italic = True
        return

    if "letterhead" in tag.get("class", []):
        for line in iter_br_separated_lines(tag):
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()
        return

    para = doc.add_paragraph()
    notes: list[str] = []
    for child in tag.children:
        process_inline(child, para, notes)
    for note in notes:
        add_note_paragraph(doc, note)


def add_list_block(doc: Document, tag: Tag, ordered: bool) -> None:
    for li in tag.find_all("li", recursive=False):
        style = "List Number" if ordered else "List Bullet"
        para = doc.add_paragraph(style=style)
        notes: list[str] = []
        for child in li.children:
            process_inline(child, para, notes)
        for note in notes:
            add_note_paragraph(doc, note)


def _appendix_lead(tag: Tag) -> bool:
    return "appendix-lead" in tag.get("class", [])


def process_body_element(doc: Document, child: Tag, intro_kind: str | None) -> None:
    if child.name in ("h1", "h2", "h3", "h4"):
        if intro_kind == "consultation" and child.name == "h1":
            return
        if child.name == "h1" and child.find(class_="changed"):
            para = doc.add_paragraph()
            notes: list[str] = []
            for sub in child.children:
                process_inline(sub, para, notes)
            for note in notes:
                add_note_paragraph(doc, note)
        else:
            add_block_heading(doc, child)
    elif child.name == "p":
        if _appendix_lead(child):
            para = doc.add_paragraph()
            run = para.add_run(child.get_text(strip=True))
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            para.paragraph_format.space_after = Pt(10)
            return
        add_paragraph_block(doc, child, intro_kind)
    elif child.name == "ul":
        add_list_block(doc, child, ordered=False)
    elif child.name == "ol":
        add_list_block(doc, child, ordered=True)
    elif child.name == "section":
        if "appendix" in child.get("class", []):
            doc.add_page_break()
        for sub in child.children:
            if isinstance(sub, Tag):
                process_body_element(doc, sub, intro_kind)


def html_to_docx(html_path: Path, doc: Document, intro_kind: str | None) -> None:
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    body = soup.find("body")
    if not body:
        raise ValueError(f"No body in {html_path}")

    for child in body.children:
        if isinstance(child, Tag):
            process_body_element(doc, child, intro_kind)


def build_comparative_analysis() -> Document:
    doc = Document()
    set_document_defaults(doc)
    add_title_page(
        doc,
        "Job Description Comparison",
    )
    html_to_docx(
        SCRIPT_DIR / "Facilities-Coordinator-JD-with-changes.html",
        doc,
        intro_kind="facilities",
    )
    add_footer(doc, "Prepared for: Facilities Coordinator (FC)")
    return doc


def build_consultation_notes() -> Document:
    doc = Document()
    set_document_defaults(doc)
    add_title_page(
        doc,
        "Consultation Response Notes: Meeting of 18 May 2026",
        subtitle="ROH follow-up letter of 20 May 2026",
    )
    html_to_docx(
        SCRIPT_DIR / "Consultation-Letter-annotated.html",
        doc,
        intro_kind="consultation",
    )
    add_footer(doc, "Prepared for: Facilities Coordinator (FC)")
    return doc


def convert_to_pdf(docx_path: Path) -> Path:
    if not SOFFICE.is_file():
        raise RuntimeError("LibreOffice not found; cannot create PDF")
    out_dir = docx_path.parent
    cmd = [
        str(SOFFICE),
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(docx_path),
    ]
    subprocess.run(cmd, check=True, capture_output=True, text=True)
    pdf_path = out_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.is_file() or pdf_path.stat().st_size < 1000:
        raise RuntimeError(f"PDF missing or too small: {pdf_path}")
    return pdf_path


def write_outputs(doc: Document, stem: str) -> list[Path]:
    written: list[Path] = []
    for out_dir in OUTPUT_DIRS:
        out_dir.mkdir(parents=True, exist_ok=True)
        docx_path = out_dir / f"{stem}.docx"
        doc.save(docx_path)
        written.append(docx_path)
        pdf_path = convert_to_pdf(docx_path)
        written.append(pdf_path)
    return written


def main() -> int:
    outputs: list[Path] = []
    comparative = build_comparative_analysis()
    consultation = build_consultation_notes()
    outputs.extend(write_outputs(comparative, "Job Description Comparison"))
    outputs.extend(write_outputs(consultation, "Consultation-Response-Notes"))
    print("Generated:")
    for p in outputs:
        size = p.stat().st_size
        print(f"  {p} ({size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
